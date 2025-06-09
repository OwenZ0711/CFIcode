# -*- coding: utf-8 -*-
"""
Created on Wed May  7 16:52:15 2025

@author: zhangziyao
"""
from datetime import date, timedelta
import matplotlib.pyplot as plt
from matplotlib.patches import Patch
import openpyxl
import os
import pandas as pd
from sqlalchemy import create_engine
from toolmodules.modules_constvars import *
from toolmodules.modules import *
import dataframe_image as dfi
import io,libconf
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import itertools

# 加入版本，t0, making置信区间
# 把切版本的拎出来看一下？


def wechat_bot_markdown(content, type_api='check'):
    url = 'https://qyapi.weixin.qq.com/cgi-bin/webhook/send?key=7a3ca5b1-c492-4aa7-993c-232ebe238f87'

    data = {
        "msgtype": "markdown",
        "markdown": {
            "content": content
        }
    }
    a = requests.post(url, json=data)
    print(a)


def wechat_bot_file(file_path, type_api='analysis-report'):
    if WECHAT_BOT_KEY_DICT.get(type_api, None) is not None:
        upload_url = f"https://qyapi.weixin.qq.com/cgi-bin/webhook/upload_media?key={WECHAT_BOT_KEY_DICT[type_api]}&type=file"
        msg_url = f"https://qyapi.weixin.qq.com/cgi-bin/webhook/send?key={WECHAT_BOT_KEY_DICT[type_api]}&type=file"
    else:
        raise 'ValueError'

    if not isinstance(file_path, list):
        file_path = [file_path]

    for file_path_i in file_path:
        data = {"file": open(file_path_i, "rb")}
        response = requests.post(url=upload_url, files=data)
        json_resp = response.json()
        media_id = json_resp["media_id"]

        msg = {
            "msgtype": "file",
            "file": {
                "media_id": media_id
            },
            "mentioned_list": [],
        }
        r = requests.post(url=msg_url, json=msg)
        print(r)

def wechat_bot_msg_check(msg, mention=None):
    if mention is None:
        mention = []
    url = 'https://qyapi.weixin.qq.com/cgi-bin/webhook/send?key=7a3ca5b1-c492-4aa7-993c-232ebe238f87'
    # 企业微信建群后在群里创建机器人，将页面的hook地址复制过来
    data = {
        "msgtype": "text",
        "text": {
            "content": msg,
            "mentioned_list": mention,
        }
    }
    a = requests.post(url, json=data)
    print(a)

class GeneratePDF():
    def __init__(self, curdate, pdf_path, max_hw=0.618):
        self.curdate = curdate
        self.pdf_path = pdf_path
        self.max_hw = max_hw
        if not os.path.exists(self.pdf_path):
            os.makedirs(self.pdf_path)

    def generate_pdf(self, pdf_filename, imgs_list):
        imgs_list_exist = []
        imgs_list_exist_not = []
        for il in imgs_list:
            if os.path.exists(il):
                imgs_list_exist.append(il)
            else:
                try:
                    imgs_list_exist_not.append(il.split('/')[-1])
                except:
                    pass
        try:
            msg = f'{pdf_filename}-[{len(imgs_list_exist)}]/[{len(imgs_list)}], ' \
                  f'不存在如下({datetime.datetime.now().strftime("%H%M%S")})：\n\t' + '\n\t'.join(imgs_list_exist_not)

            # wechat_bot_msg_check(msg)
        except:
            pass

        width, height, hw_ratio, img_size_dict = 0, 0, 0, {}
        for img in imgs_list_exist:
            with Image.open(img) as img_file:
                w, h = img_file.size
            width, height = max(w, width), max(h, height)
            hw_ratio = min(max(h / w, hw_ratio), self.max_hw)
            img_size_dict[img] = [w, h, h / w]

        if height / width < hw_ratio:
            width, height = width, int(width * hw_ratio)
        else:
            width, height = int(height / hw_ratio), height

        hw_ratio_pdf = height / width
        width += 10
        height += 10

        pdf = FPDF(unit="pt", format=[width, height])
        for page in imgs_list_exist:
            pdf.add_page()

            w, h, hw_ratio = img_size_dict[page]
            if hw_ratio <= hw_ratio_pdf:
                w, h = width - 10, int(h * (width - 10) / w)
            else:
                w, h = int(w * (height - 10) / h), height - 10

            x = int((width - w) / 2)
            y = int((height - h) / 2)
            pdf.image(page, x, y, w, h)

        pdf.output(self.pdf_path + pdf_filename, "F")


SINGLE_COLUMNS = ['SZ_t2o(mean)', 'SZ_t2o(median)', 'SZ_t2o(75%)',
   'SH_t2o(mean)', 'SH_t2o(median)', 'SH_t2o(75%)', 'engine (SH)',
   'strategy (SH)', 'send (SH)', 'engine (SZ)', 'strategy (SZ)',
   'send (SZ)']
Orig_Columns = ['md engine delay (SZ)', 'strategy delay (SZ)', 'send_delay (SZ)','SZ_tick2order(mean)',
                'SZ_tick2order(median)', 'SZ_tick2order(75%)','md engine delay (SH)', 'strategy delay (SH)',
                'send_delay (SH)','SH_tick2order(mean)', 'SH_tick2order(median)', 'SH_tick2order(75%)']

Single_Rename_Col = {'SZ_tick2order(mean)':'SZ_t2o(mean)', 'SZ_tick2order(median)':'SZ_t2o(median)', 'SZ_tick2order(75%)':'SZ_t2o(75%)','SH_tick2order(mean)':'SH_t2o(mean)',
'SH_tick2order(median)':'SH_t2o(median)', 'SH_tick2order(75%)':'SH_t2o(75%)','md engine delay (SH)':'engine (SH)', 'strategy delay (SH)':'strategy (SH)',
'send_delay (SH)':'send (SH)','md engine delay (SZ)':'engine (SZ)', 'strategy delay (SZ)':'strategy (SZ)', 'send_delay (SZ)':'send (SZ)'}

DUAL_COLUMNS = ['t2o(mean)', 't2o(median)', 't2o(75%)', 'engine','strategy', 'send']
CANCEL_COLUMNS = ['撤单股数比(SZ)', '撤单金额比(SZ)','撤单股数比(SH)', '撤单金额比(SH)','撤单股数比','撤单金额比']

TARGET_SINGLE_COLUMNS = ['SZ_t2o(mean)_agg', 'SZ_t2o(median)_agg', 'SZ_t2o(75%)_agg',
   'SH_t2o(mean)_agg', 'SH_t2o(median)_agg', 'SH_t2o(75%)_agg', 'engine (SH)_agg',
   'strategy (SH)_agg', 'send (SH)_agg', 'engine (SZ)_agg', 'strategy (SZ)_agg',
   'send (SZ)_agg']
TARGET_DUAL_COLUMNS = ['t2o(mean)_agg', 't2o(median)_agg', 't2o(75%)_agg', 'engine_agg','strategy_agg', 'send_agg']
TARGET_CANCEL_COLUMNS = ['撤单股数比(SZ)_agg', '撤单金额比(SZ)_agg','撤单股数比(SH)_agg', '撤单金额比(SH)_agg','撤单股数比_agg','撤单金额比_agg']
Long_Short_prod = [ 'MSLS1','SHORTMSLS1','MSLS2','SHORTMSLS2']
Own_Prod_List = ['HXDC21','WPGZQ1','DC19','DC19B']



def find_t0_prod_list(curr_date):
    with io.open(rf"V:\StockData\d0_file\adjust_cfg_paras\paras_config\{curr_date}_paras_config.cfg") as f:
        config = libconf.load(f)
    return list(set((config['open_t0_mode_sh'][0]['target_list'] + config['open_t0_mode_sz'][0]['target_list'])))

today = date.today()
today_str = today.strftime('%Y%m%d')
# today_str = '20250604'
print(today_str)
ten_days_ago = get_predate(today_str,n = 10)
Config_T0_All = find_t0_prod_list(today_str)


def get_pta_report(start_date, end_date):
    '''
    外部 莫改
    pta源头数据读取
    '''
    engine = create_engine("mssql+pymssql://ht:huangtao@dbs.cfi/DataSupply")
    data_sql = '''
        select * from
        dbo.stock_pta_report
        where date between '%s' and '%s'
    '''

    data_str = data_sql % (start_date, end_date)
    df = pd.read_sql(data_str, engine)
    return df


def process_daily_average_change(original_df, cancel, dual):
    if cancel == True:
        target_col = CANCEL_COLUMNS.copy()
    else:
        if dual == True:
            target_col = DUAL_COLUMNS.copy()
        else:
            target_col = SINGLE_COLUMNS.copy()
    temp_df = original_df.copy().sort_values(by='date',ascending=True)
    del temp_df['date']
    for col in target_col:
        temp_df[col+'_incre'] = temp_df.groupby(['production', 'center'])[col].diff()
    increment_means = temp_df.groupby(['production', 'center'])[[f'{col}_incre' for col in target_col]].mean().reset_index()
    increment_count = temp_df.groupby(['production', 'center'])[[f'{col}_incre' for col in target_col]].apply(lambda x: (x > 0).sum()).reset_index()
    return increment_means, increment_count   


def check_own_acc(df):
    df["自营"] = df["production"].isin(Production_OwnList_Swap + Long_Short_prod + Own_Prod_List)

def check_specific_own_acc(input_df,prod_type):
    """
    helper
    筛选自营subset
    """
    df = input_df.copy()
    df["自营"] = df["production"].isin(prod_type)
    return df

def find_colo_shsz(colo):
    '''
    helper
    用colo识别深圳上海中心
    '''
    if colo[:7] == "cfipasz":
        return "sz"
    if colo[:7] == "cfipash":
        return "sh"
    start_char = colo.find('-') + 1
    end_char = colo.rfind("-")
    return colo[start_char:end_char] if colo[start_char:end_char] == "sh" else "sz"

def make_cols_into_str(row,col):
    '''
    apply function
    表格中确定单元格格式
    '''
    return f'{row[col][0]}, {row[col][1]}'

def make_cols_into_1(row, col):
    '''
    apply function
    表格中确定单元格4位小数方便视图
    '''
    return f"{format(row[col+'_orig'],'.4f')}- {format(row[col+'_10日'],'.4f')}"

def get_difference_ratio(row,column, past_column):
    '''
    apply function
    当前值与10天平局差值百分比，过去为0则显示100异值
    '''
    if (row[column]-row[past_column])==0:
        return 0
    elif row[past_column] == 0:
        return 100
    else:
        return (row[column]-row[past_column])/row[past_column]


def find_diff_t2o(df,single=1):
    """
    计算与10日平均差值并创建列
    """
    column_list = DUAL_COLUMNS.copy()
    if single == 1:
        column_list = SINGLE_COLUMNS.copy()
    for column in column_list:
        new_column = column+"_diff"
        past_column = column+"_10日"
        df[new_column] = df.apply(get_difference_ratio,args=(column,past_column,),axis=1)
    return df


def only_should_contain1(row, column):
    '''
    双中心产品t2o会有很多nan记录，只保留本中心的记录, 格式1
    '''
    if pd.notna(row["SH_"+column]) and pd.notna(row["SZ_"+column]):
        assert print(f"something is wrong with {row['production']} {column} column,should not have both SH and SZ")
    else:
        return row["SH_"+column] if pd.notna(row["SH_"+column]) else row["SZ_"+column]

def only_should_contain2(row, column):
    '''
    双中心产品t2o会有很多nan记录，只保留本中心的记录, 格式2
    '''
    if pd.notna(row[column+" (SH)"]) and pd.notna(row[column+" (SZ)"]):
        assert print(f"something is wrong with {row['production']},should not have both SH and SZ")
    else:
        return row[column+" (SH)"] if pd.notna(row[column+" (SH)"]) else row[column+" (SZ)"]

def get_clean_df(df):
    '''
    helper 把双中心的nan值清理掉
    '''
    column_list1 = ['t2o(mean)', 't2o(median)', 't2o(75%)']
    column_list2=['engine', 'strategy','send']
    column_del = []
    for column in column_list1:
        print(column)
        df[column] = df.apply(only_should_contain1,args=(column,),axis=1)
        column_del.append("SH_"+column)
        column_del.append("SZ_"+column)
    for column in column_list2:
        print(column)
        df[column] = df.apply(only_should_contain2,args=(column,),axis=1)
        column_del.append(column+" (SH)")
        column_del.append(column+ " (SZ)")
    for column in column_del:
        del df[column]
    return df
    

def combine_cols(df,single=1):
    '''
    创建_agg列（作图用），每个entry存成list模式并且保留4位小数
    '''
    df_copy = df.copy()
    if single == 1:
        type_cols = SINGLE_COLUMNS
    else:
        type_cols = DUAL_COLUMNS
    for col in type_cols:
        df_copy[col+'_orig'] = df_copy[col]
        df_copy[col] = df_copy.apply(make_cols_into_1,args=(col,),axis=1)
    return df_copy

def filter_condition(row,type_cols):
    '''
    apply func
    筛选问题产品的条件
    '''
    if row["production"] in Production_OwnList_Swap + Long_Short_prod + Own_Prod_List:
        return True
    # if "cms" in row["colo"] and "5" in row["colo"]:#指定
    #     return True
    for col in type_cols:
        if row[col+'_orig'] == 0 and row[col+"_incre_count"] > 7:
            return True
        if row[col+'_orig'] != 0 and (row[col+"_incre"]/row[col+'_orig']) > 0.1 and row[col + '_incre_count'] > 5:
            return True
        if row[col+"_diff"] > 1:
            return True
    return False

def filter_df(orig_df, single=1):
    '''
    筛选符合问题的产品 t2o
    '''
    if single == 1:
        type_cols = SINGLE_COLUMNS
    else:
        type_cols = DUAL_COLUMNS
    df = orig_df.copy()
    mask_df = df.apply(filter_condition,args=(type_cols,),axis = 1)
    df = df.loc[mask_df]
    return df

def color_values(target_col,df):
    '''
    apply func
    根据筛选条件增加颜色
    '''
    suffix = target_col.name
    filter1 = df[suffix+"_diff"] > 1
    filter2 = ((df[suffix+'_orig'] > 0) & (df[suffix + "_incre"]/df[suffix+'_orig'] > 0.1) & (df[suffix + '_incre_count'] > 5)) | (df[suffix + '_incre_count'] > 7)
    filter4 = df[suffix+'_diff'] < -0.5
    res = []
    for idx in target_col.index:
        if filter2[idx]: #增量异常
            res.append('background-color: lime')
        # elif filter3[idx]:
        #     res.append('background-color: #FF9999')
        elif filter1[idx]: #当天过高
            res.append("background-color: lightblue")
        # elif filter4[idx]:
        #     res.append("background-color: #00FFFF")
        else:
            res.append("")
    return res
    
    
    # if target_col[siffix+"_diff"] > 0.3:
    #     return 'background-color: red'  # Pinkish for negative
    # if target_col[suffix] > (target_col[suffix+"_10日"] + 2*target_col[suffix+"_std"]):
    #     return "background-color: blue"
    # if (target_col[suffix] - target_col[suffix+"_10日"]) > 0.5:
    #     return 'background-color: yellow'
    # return ''

def process_df_image(original_df,single=1):
    '''
    生成加了颜色的表格
    '''
    temp_target_columns = []
    if single == 1:
        temp_target_columns = SINGLE_COLUMNS.copy()
    else:
        temp_target_columns = DUAL_COLUMNS.copy()

    colored_original_df = original_df.fillna("0").style \
        .set_table_styles([
        dict(selector='td', props=[('border', '1px solid black')]),
        dict(selector='caption', props=[('padding', '1px'), ('margin', '1px 0'), ('text-align', 'center')])
        ]).apply(color_values, args=(original_df,), subset=temp_target_columns)  \
        .apply(color_cancel_type,  subset=['production'],axis = 0) \
        .set_properties(**{'text-align': 'center', 'border': '1px solid black'}) \
        .hide([col for col in original_df.columns if col not in temp_target_columns+["production",'colo','指数']], axis=1) \
        .set_caption(f'''
        <div style="font-size: 10px">
            <strong>{today_str}t2o筛选条件</strong><br>
            <div style="font-size: 6px; display: flex; align-items: center; margin-bottom: 1px;">
                与10日均差100%： <span style="background-color: lightblue; width: 10px; height: 4px; display: inline-block; margin-right: 3px;"></span> ;  增量异常： <span style="background-color: lime; width: 10px; height: 4px; display: inline-block; margin-right: 3px;"></span>;   单位：秒
            </div>
        </div>
        ''')
        
    return colored_original_df

def find_colo(colo_name):
    if colo_name[:5] == "cfipa":
        return "pa"+'-'+colo_name[7:]
    else:
        i_front = colo_name.find('-')
        i_rear = colo_name.rfind('-')
        return colo_name[:i_front] + colo_name[i_rear:]

def find_broker(colo_name):
    if colo_name[:5] == "cfipa":
        return "pa"
    else:
        i_front = colo_name.find('-')
        return colo_name[:i_front]


def return_df_dict(orig):
    '''
    每10行输出成一个df，限制图片长度
    '''
    # process colo
    df = orig.copy()
    df['broker'] = df['colo'].apply(lambda x: find_broker(x))
    df['simp_colo'] = df['colo'].apply(lambda x: find_colo(x))
    all_colo = list(set(df['simp_colo']))
    all_broker_dict = {}
    for i in range(len(all_colo)):
        if find_broker(all_colo[i]) in all_broker_dict.keys():
            all_broker_dict[find_broker(all_colo[i])].append(all_colo[i])
        else:
            all_broker_dict[find_broker(all_colo[i])] = [all_colo[i]]
    df_dict = {}
    for key in all_broker_dict.keys():
        broker_df_list = []
        for val in all_broker_dict[key]:
            colo_df =df[df['simp_colo'] == val].sort_values('production')
            broker_df_list.append(colo_df)
        df_dict[key] = broker_df_list
    
    return df_dict



def plot_t2o_curves(date,exch,target_products,target_cols, x, y):
    assert (exch == "sh" or exch == 'sz'), "center is wrong, has to be sh or sz"
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.rcParams['axes.unicode_minus'] = False
    temp_end_date = get_predate(date,n = 10)
    original_df_graph = get_pta_report(temp_end_date, date)
    for target_prod in target_products:
        if target_prod not in list(original_df_graph['production']):
            print(target_prod)
            assert f"{target_prod} is not a valid production code"
    for target_single_col in target_cols:
        if target_single_col not in original_df_graph.columns:
            print("Below are valid columns")
            print(original_df_graph.columns)
            assert f"{target_single_col} is not a valid column name"
    df = original_df_graph.loc[original_df_graph['production'].isin(target_products)]
    df = df.sort_values(by='date')
    df = df[['production','colo','date']+target_cols]
    df['center'] = df['colo'].apply(find_colo_shsz)
    df = df[df['center'] == exch]
    del df["center"], df['colo']

    fig, axes = plt.subplots(x, y, figsize=(9,12), sharex=True)
    axes = axes.flatten()
    # Plot each cancellation column in its own subplot
    for i, col in enumerate(target_cols):
        ax = axes[i]
        # Pivot data: rows are dates, columns are productions
        pivot_df = df.pivot(index='date', columns='production', values=col)
        # min_val = 30
        # for prod in pivot_df.columns:
        #     if len(pivot_df[prod] == 0):
        #         assert f"{prod}没有10天数据，需要检查"
        #         pass
        #     else:
        #         day_10_col = []
        #         for min_i in range(len(pivot_df[prod])):
        #             if pivot_df[prod][min_i]:
        #                 day_10_col.append(pivot_df[prod][min_i])
        #         if prod_min < min_val:
        #             min_val = prod_min
        # max_val = pivot_df.max().max()
        ax.plot(pivot_df, label=pivot_df.columns,marker='o')
        for idx, production in enumerate(pivot_df.columns):
            # Add number at the end of each line
            mean = pivot_df[production].mean()
            line = ax.lines[idx]  # Get the line object for this production
            color = line.get_color()
            ax.axhline(y=mean, color=color, linestyle='--')
            last_date = pivot_df.index[-1]
            last_value = pivot_df[production].iloc[-1]
            text = f"{last_value:.2f}"
            ax.annotate(text, 
                       xy=(last_date, last_value), 
                       xytext=(5, 0),  # Offset 5 points to the right
                       textcoords='offset points', 
                       fontsize=10, 
                       ha='left', 
                       va='center')
        ax.grid(True)
        ax.set_title(col, fontsize=16)
        ax.set_xlabel('Date', fontsize=16)
        ax.set_ylabel(col, fontsize=16)
        ax.tick_params(axis='both', labelsize=16)
        #ax.set_ylim(0, 100)
    handles,labels = axes[0].get_legend_handles_labels()
    # Add a single legend for all subplots
    fig.legend(handles, labels, loc='upper right', bbox_to_anchor=(0.7, 1.05), ncol=4,title='Production')
    fig.suptitle(f'手动画图分析 {today_str} exchange:{exch}', fontsize=20, y=1.08)
    # Format x-axis dates and adjust layout
    fig.autofmt_xdate()
    plt.tight_layout()
    plt.savefig(rf'V:\StockData\Trading\owen_code\pta_result\{date}\{date}_own_t2o_curve_{exch}.png', dpi=300, bbox_inches='tight', transparent=False)
    plt.show()



'''
以下为发撤单函数，其中包含了绘图自营产品
'''
def make_cols_into_1_cancel(row, col):
    '''
    apply function
    表格中确定单元格2位小数方便视图
    '''
    return f"{format(row[col+'_orig'],'.1f')}- {format(row[col+'_10日'],'.1f')}"

def combine_cancel_cols(df):
    print("combining cancel")
    df_copy = df.copy()
    type_cols = CANCEL_COLUMNS
    for col in type_cols:
        print(col)
        df_copy[col+"_orig"] = df_copy[col]
        df_copy[col] = df_copy.apply(make_cols_into_1_cancel,args=(col,),axis=1)
    return df_copy


def filter_cancel_condition(row,type_cols):
    if row["production"] in Production_OwnList_Swap + Long_Short_prod + Own_Prod_List:
        return True
    for col in type_cols:
        # if "cms" in row["colo"] and "5" in row["colo"]: #指定
        #     return True
        filter_incre = False
        if row[col+'_orig'] == 0 and row[col+"_incre_count"] > 7:
            filter_incre = True
        if row[col+"_orig"] > 20 and row[col+"_orig"] != 0 and (row[col+"_incre"]/row[col+"_orig"]) > 0.1 and row[col + '_incre_count'] > 5:
            filter_incre = True
        if row[col + "_diff"] < -0.5 and row[col+'_10日'] > 10 and (row['production'] in Config_T0_All):
            filter_incre = True
        if (row[col+"_diff"] > 0.3 and row[col+"_orig"]> 20) or row[col+"_orig"]> 40:
            filter_incre = True
        if filter_incre and (row[col+"_orig"] > 20 or row[col+"_orig"] < 10):
            return True

    return False

def filter_cancel_df(orig_df):
    type_cols = CANCEL_COLUMNS
    df = orig_df.copy()
    mask_df = df.apply(filter_cancel_condition,args=(type_cols,),axis = 1)
    df = df.loc[mask_df]
    return df
    

def filter_own_cancel_condition(row,type_cols):
    for col in type_cols:
        filter_incre = False
        # if "cms" in row["colo"] and "5" in row["colo"]: #指定
        #     return True
        if row[col+'_orig'] == 0 and row[col+"_incre_count"] > 7:
            filter_incre = True
        if row[col+"_orig"] > 20 and row[col+"_orig"] != 0 and (row[col+"_incre"]/row[col+"_orig"]) > 0.1 and row[col + '_incre_count'] > 5:
            filter_incre = True
        if row[col + "_diff"] < -0.5 and row[col+'_10日'] > 10 and (row['production'] in Config_T0_All):
            filter_incre = True
        if (row[col+"_diff"] > 0.3 and row[col+"_orig"]> 20) or row[col+"_orig"]> 40:
            filter_incre = True
        if filter_incre and (row[col+"_orig"] > 20 or row[col+"_orig"] < 10):
            return True

    return False

def filter_owen_cancel_df(orig_df):
    type_cols = CANCEL_COLUMNS
    df = orig_df.copy()
    mask_df = df.apply(filter_own_cancel_condition,args=(type_cols,),axis = 1)
    df = df.loc[mask_df]
    return df


def color_cancel_values(target_col,df):
    suffix = target_col.name
    filter_out = (df[suffix+'_orig'] < 20) & (df[suffix+'_orig'] > 10)
    filter1 = (df[suffix+"_diff"] > 0.3) & (df[suffix+'_orig'] > 20)
    filter2 = df[suffix+"_orig"] > 40
    filter3 = ((df[suffix+'_orig'] > 20) & (df[suffix+'_orig'] != 0) & (df[suffix + "_incre"]/df[suffix+'_orig'] > 0.1) & (df[suffix + '_incre_count'] > 5)) | (df[suffix + '_incre_count'] > 7)
    # filter4 = ((df[suffix+"_orig"] < 10) | (df[suffix+'_diff'] < -0.5)) & df['production'].isin(Config_T0_All)
    filter4 =  (df[suffix+'_diff'] < -0.5) & (df[suffix+'_10日'] > 10) & df['production'].isin(Config_T0_All)
    res = []
    for idx in target_col.index:
        if filter_out[idx]:
            res.append("")
        elif filter3[idx]:
            res.append("background-color: lime")
        elif filter2[idx]:
            res.append("background-color: #FF9999")
        elif filter1[idx]:
            res.append('background-color: lightblue')
        elif filter4[idx]:
            res.append("background-color: #00FFFF")
        else:
            res.append("")
    return res

def color_cancel_type(target_col):
    res = ['background-color: yellow' if x in Production_OwnList_Swap+Long_Short_prod+Own_Prod_List else "" for x in target_col]
    return res
    

def process_cancel_image(original_df,single=1):
    print('making table')
    temp_target_columns = CANCEL_COLUMNS
    colored_original_df = original_df.fillna("0").style \
        .set_table_styles([
        dict(selector='td', props=[('border', '1px solid black')]),
        dict(selector='caption', props=[('padding', '1px'), ('margin', '1px 0'), ('text-align', 'center')])
        ]).apply(color_cancel_values, args=(original_df,), subset=temp_target_columns)  \
        .apply(color_cancel_type,  subset=['production'],axis = 0) \
        .set_properties(**{'text-align': 'center', 'border': '1px solid black'}) \
        .hide([col for col in original_df.columns if col not in temp_target_columns+["production",'colo',"指数"]], axis=1) \
        .set_caption(f'''
        <div style="font-size: 10px">
            <strong>{today_str}撤单筛选条件</strong><br>
            <div style="font-size: 6px; display: flex; align-items: center; margin-bottom: 1px;">
                与10日均差30%： <span style="background-color: lightblue; width: 10px; height: 4px; display: inline-block; margin-right: 3px;"></span> ;  大于40%： <span style="background-color: #FF9999; width: 10px; height: 4px; display: inline-block; margin-right: 3px;"></span> ;   增量异常： <span style="background-color: lime; width: 10px; height: 4px; display: inline-block; margin-right: 3px;"></span> ;   自营产品： <span style="background-color: yellow; width: 10px; height: 4px; display: inline-block; margin-right: 3px;"></span> ; 骤降： <span style="background-color: #00FFFF; width: 10px; height: 4px; display: inline-block; margin-right: 3px;"></span> ;   单位：百分比
            </div>
        </div>
        ''')
        #.set_properties(subset=temp_target_columns, **{'width': '20px'}) \
        #.hide(hide_columns)
        #.background_gradient(cmap='RdYlGn', subset=['tick2order(mean)_diff', 'tick2order(median)_diff',
        #'tick2order(75%)_diff', 'md engine delay_diff', 'strategy delay_diff',
        #'send_delay_diff'], vmin=0, vmax=1)
        
    return colored_original_df

    
def find_diff_cancel(df,single=1):
    """
    计算与10日平均差值并创建列
    """
    column_list = CANCEL_COLUMNS.copy()
    for column in column_list:
        new_column = column+"_diff"
        past_column = column+"_10日"
        df[new_column] = df.apply(get_difference_ratio,args=(column,past_column,),axis=1)
    return df


def combine_own_cancel_columns(original_df):
    original_df = original_df.groupby('production')

    
def plot_cancel_curves(cancel_df_graph,date,label):
    plt.rcParams['font.sans-serif'] = ['SimHei']
    plt.rcParams['axes.unicode_minus'] = False
    df = cancel_df_graph.copy()
    
    # Sort the DataFrame by 'date'
    df = df.sort_values(by='date')
    if label == "LS":
        target_col_graph = ['撤单金额比(SZ)']
        fig, axes = plt.subplots(1, 1, figsize=(12,6), sharex=True)
    else:
        target_col_graph = ['撤单金额比(SZ)', '撤单金额比(SH)']
        fig, axes = plt.subplots(2, 1, figsize=(12,12), sharex=True)
        axes = axes.flatten()
    # Plot each cancellation column in its own subplot
    for i, col in enumerate(target_col_graph):
        if label == "LS":
            ax = axes
        else:
            ax = axes[i]
        # Pivot data: rows are dates, columns are productions,sz中心看sz撤单率，sh中心看sh撤单率
        if col == '撤单金额比(SZ)':
            new_df = df[df['center'] == 'sz']
        else:
            new_df = df[df['center'] == 'sh']
        pivot_df = new_df.pivot(index='date', columns='production', values=col)
        ax.plot(pivot_df, label=pivot_df.columns, marker='o')
        mean_labels = []
        for idx, production in enumerate(pivot_df.columns):
            # Add number at the end of each line
            mean = pivot_df[production].mean()
            line = ax.lines[idx]  # Get the line object for this production
            color = line.get_color()
            ax.axhline(y=mean, color=color, linestyle='--')
            last_date = pivot_df.index[-1]
            last_value = pivot_df[production].iloc[-1]
            text = f"{last_value:.2f}"
            ax.annotate(text, 
                       xy=(last_date, last_value), 
                       xytext=(5, 0),  # Offset 5 points to the right
                       textcoords='offset points', 
                       fontsize=10, 
                       ha='left', 
                       va='center')
            mean_labels.append(f"{production}: 均值: {mean:.2f} \n        当日: {last_value}")
        dummy_handles = [Patch(color='none', label=label) for label in mean_labels]
        ax.legend(handles=dummy_handles, title="中心汇总", loc='upper right', bbox_to_anchor=(0, 1), fontsize=9, ncol=1)
        ax.grid(True)
        ax.set_title(f"{date}:{col}", fontsize=16)
        ax.set_xlabel('Date', fontsize=16)
        ax.set_ylabel(col, fontsize=16)
        ax.tick_params(axis='both', labelsize=16)
        #ax.set_ylim(0, 100)
    if label == 'LS':
        handles,labels = axes.get_legend_handles_labels()
    else:
        handles,labels = axes[0].get_legend_handles_labels()
    # Add a single legend for all subplots
    this_type = '自营' if label == "own" else "LongShort"
    fig.legend(handles, labels, loc='upper right', bbox_to_anchor=(0.7, 1.05), ncol=4,title='Production')
    fig.suptitle(f'自营撤单汇总 {date} type:{this_type}', fontsize=20, y=1.08)
    # Format x-axis dates and adjust layout
    fig.autofmt_xdate()
    plt.tight_layout()
    plt.savefig(rf'V:\StockData\Trading\owen_code\pta_result\{date}\{date}_{label}_cancel_curve.png', dpi=300, bbox_inches='tight', transparent=False)
    plt.show()
    wechat_bot_image(rf'V:\StockData\Trading\owen_code\pta_result\{date}\{date}_{label}_cancel_curve.png',type_api='analysis-report')
    
    
def get_top_cancel_ticker_sum(input_str):
    return sum(float(x) for x in input_str.split(',') if float(x) > 1)
def get_top_cancel_ticker_sum_sh(input_series):
    res_list = []
    cancel_ticker = input_series['撤单TpCL'].split(',')
    cancel_percentage = input_series['撤单TpCrL'].split(',')
    if len(cancel_ticker) != len(cancel_percentage):
        assert "length not match gr df"
    for i in range(len(cancel_ticker)):
        if cancel_ticker[i][:2] == "SH" and float(cancel_percentage[i]) > 1:
            res_list.append(float(cancel_percentage[i]))
    return sum(res_list)
    
def get_top_cancel_ticker_sum_sz(input_series):
    res_list = []
    cancel_ticker = input_series['撤单TpCL'].split(',')
    cancel_percentage = input_series['撤单TpCrL'].split(',')
    if len(cancel_ticker) != len(cancel_percentage):
        assert "length not match gr df"
    for i in range(len(cancel_ticker)):
        if cancel_ticker[i][:2] == "SZ" and float(cancel_percentage[i]) > 1:
            res_list.append(float(cancel_percentage[i]))
    return sum(res_list)

def get_broker_t2o_error_count(type_dict):
    error_list_all = []
    for broker in type_dict.keys():
        error_colo_msg_lst = []
        broker_error_num = 0
        
        for colo_df in type_dict[broker]:
            error_prod_num = len(set(colo_df['production']))
            colo_name = list(colo_df['simp_colo'])[0]
            if error_prod_num != 0:
                colo_prod_problem = {}
                for prd_name in set(colo_df['production']):
                    colo_prd_error_msg_series = colo_df[colo_df['production'] == prd_name].apply(get_t2o_problem_type, axis = 1)
                    colo_prd_error_msg_lst_of_lst = [msg.split(',') for msg in colo_prd_error_msg_series]
                    colo_prd_error_msg_lst = set([error_item for sublist in colo_prd_error_msg_lst_of_lst for error_item in sublist])
                    colo_prd_error_msg = ",".join(colo_prd_error_msg_lst)
                    colo_prod_problem[prd_name] = colo_prd_error_msg
                error_colo_msg_lst.append(f"{colo_name}:\n{colo_prod_problem}\n")
            broker_error_num += error_prod_num
            
        if broker_error_num != 0:
            error_colo_msg = "".join(error_colo_msg_lst)
            error_broker_msg = f"{broker} 有 {broker_error_num}个问题: {error_colo_msg}"
            error_list_all.append(error_broker_msg)
    if len(error_list_all) == 0:
        return "无错误"
    else:
        error_msg = "\n".join(error_list_all)
        return error_msg

def get_t2o_problem_type(row):
    type_cols = DUAL_COLUMNS
    errors = []
    for col in type_cols:
        if row[col+'_orig'] == 0 and row[col+"_incre_count"] > 7:
            errors.append("增量异常")
        if row[col+'_orig'] != 0 and (row[col+"_incre"]/row[col+'_orig']) > 0.1 and row[col + '_incre_count'] > 5:
            errors.append("增量异常")
        if row[col+"_diff"] > 1:
            errors.append("延迟过高")
    error_str = ",".join(set(errors))

    return error_str

def get_broker_cancel_error_count(type_dict):
    error_list_all = []
    for broker in type_dict.keys():
        error_colo_msg_lst = []
        broker_error_num = 0
        
        for colo_df in type_dict[broker]:
            error_prod_num = len(set(colo_df['production']))
            colo_name = list(colo_df['simp_colo'])[0]
            if error_prod_num != 0:
                colo_prod_problem = {}
                for prd_name in set(colo_df['production']):
                    colo_prd_error_msg_series = colo_df[colo_df['production'] == prd_name].apply(get_cancel_problem_type, axis = 1)
                    colo_prd_error_msg_lst_of_lst = [msg.split(',') for msg in colo_prd_error_msg_series]
                    colo_prd_error_msg_lst = set([error_item for sublist in colo_prd_error_msg_lst_of_lst for error_item in sublist])
                    colo_prd_error_msg = ",".join(colo_prd_error_msg_lst)
                    colo_prod_problem[prd_name] = colo_prd_error_msg
                error_colo_msg_lst.append(f"{colo_name}:\n{colo_prod_problem}\n")
            broker_error_num += error_prod_num
            
        if broker_error_num != 0:
            error_colo_msg = "".join(error_colo_msg_lst)
            error_broker_msg = f"{broker} 有 {broker_error_num}个问题: {error_colo_msg}"
            error_list_all.append(error_broker_msg)
    if len(error_list_all) == 0:
        return "无错误"
    else:
        error_msg = "\n".join(error_list_all)
        return error_msg


def get_cancel_problem_type(row):
    type_cols = CANCEL_COLUMNS
    errors = []
    for col in type_cols:
        if row[col+"_orig"] < 20 and row[col+"_orig"] > 10:
            continue
        if row[col+'_orig'] == 0 and row[col+"_incre_count"] > 7:
            errors.append("增量异常")
        if row[col+"_orig"] > 20 and row[col+"_orig"] != 0 and (row[col+"_incre"]/row[col+"_orig"]) > 0.1 and row[col + '_incre_count'] > 5:
            errors.append("增量异常")
        if row[col + "_diff"] < -0.5 and row[col+'_10日'] > 10 and (row['production'] in Config_T0_All):
           errors.append("T0骤减")
        if (row[col+"_diff"] > 0.3 and row[col+"_orig"]> 20) or row[col+"_orig"]> 40:
            errors.append("比例过高")
    error_str = ",".join(set(errors))
    return error_str
            

def get_cancel_own_summary_msg(cancel_own, cancel_own_selected):
    
    def if_cancel_own_problem(row):
        if row["production"] in cancel_own_selected:
            return get_cancel_problem_type(row)
        else:
            return "无误"
    
    doc = Document()
    cancel_own_summary_msg_list = list(cancel_own.apply(lambda row: f"{row['production']} {row['colo']}: \n撤单股数比 {row['撤单股数比(SZ)']}; 撤单金额比 {row['撤单金额比(SZ)']}; {if_cancel_own_problem(row)}" if row['center'] == "sz" else f"{row['production']} {row['colo']}: \n撤单股数比 {row['撤单股数比(SH)']}; 撤单金额比 {row['撤单金额比(SH)']}; {if_cancel_own_problem(row)}",axis=1))  
    cancel_own_summary_msg_ok_lst = []
    cancel_own_summary_msg_error_lst = []
    for msg in cancel_own_summary_msg_list:
        if "无误" in msg:
            cancel_own_summary_msg_ok_lst.append(msg)
        else:
            cancel_own_summary_msg_error_lst.append(msg)
    cancel_own_summary_msg_ok = "\n".join(cancel_own_summary_msg_ok_lst)
    cancel_own_summary_msg_error = "\n".join(cancel_own_summary_msg_error_lst)
    
    heading = doc.add_heading(f'{today_str}自营撤单信息汇总', level=1)   
    for run in heading.runs:
        run.font.name = 'SimSun'  # Set font to SimSun for Chinese
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')  # Ensure East Asian font
        run.font.size = Pt(15)
    if len(cancel_own_summary_msg_error_lst) != 0:
        heading2 = doc.add_heading(f'问题产品', level=1)   
        for run in heading2.runs:
            run.font.name = 'SimSun'  # Set font to SimSun for Chinese
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')  # Ensure East Asian font
            run.font.size = Pt(14)
        paragraph = doc.add_paragraph(cancel_own_summary_msg_error)  
        for run in paragraph.runs:
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(11)
    if len(cancel_own_summary_msg_ok_lst) != 0:
        heading3 = doc.add_heading(f'无误产品', level=1)   
        for run in heading3.runs:
            run.font.name = 'SimSun'  # Set font to SimSun for Chinese
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')  # Ensure East Asian font
            run.font.size = Pt(14)
        paragraph2 = doc.add_paragraph(cancel_own_summary_msg_ok)  
        for run in paragraph2.runs:
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(11)
        
    doc.save(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\自营撤单信息汇总.docx")
    wechat_bot_file(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\自营撤单信息汇总.docx")
    return cancel_own_summary_msg_ok + "/n" + cancel_own_summary_msg_error
    

def error_output_creation(single,single_own,dual_dict,dual,dual_own,cancel_dict,cancel,cancel_own):
    cancel_own_selected_df = filter_owen_cancel_df(cancel_own)
    cancel_own_selected = list(cancel_own_selected_df['production'])
    single_own_selected = filter_df(single_own)
    dual_own_selected = filter_df(dual_own,single=0)
    
    # tick 2 order 只处理有问题的
    general_msg = "tick to order及撤单率问题汇总\n"
    # 自营 tick to order
    tick2order_own_msg = "当日自营tick to order问题汇总：\n"
    single_own_msg = "无错误" if len(single_own_selected) == 0 else ",".join(list(single_own_selected['production']))
    dual_own_msg = "无错误" if len(dual_own_selected) == 0 else ",".join(list(dual_own_selected['production']))
    print(tick2order_own_msg)
    print("单中心自营问题summary:\n"+single_own_msg)
    print("双中心自营问题summary:\n"+dual_own_msg)
    tick2order_single_msg = "无错误" if len(single) == 0 else ",".join(list(single['production']))
    tick2order_dual_msg = get_broker_t2o_error_count(dual_dict)
    cancel_own_summary_msg = get_cancel_own_summary_msg(cancel_own,cancel_own_selected)
    cancel_own_msg = "无错误" if len(cancel_own_selected_df) == 0 else ",".join(set(cancel_own_selected_df['production']))
    cancel_all_msg = get_broker_cancel_error_count(cancel_dict)
    print("管理单中心问题summary:\n"+tick2order_single_msg)
    print("管理双中心问题summary:\n"+tick2order_dual_msg)
    print("撤单自营报告summary:\n" + cancel_own_summary_msg)
    print("撤单自营问题summary:\n" + cancel_own_msg)
    print("撤单管理问题summary:\n" + cancel_all_msg)
    doc = Document()
    heading = doc.add_heading(f'{today_str}管理撤单问题汇总', level=1)    
    for run in heading.runs:
        run.font.name = 'SimSun'  # Set font to SimSun for Chinese
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')  # Ensure East Asian font
        run.font.size = Pt(15)
    
    paragraph = doc.add_paragraph(cancel_all_msg)  
    for run in paragraph.runs:
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.size = Pt(11)
        
    doc.save(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\管理撤单问题汇总.docx")
    wechat_bot_file(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\管理撤单问题汇总.docx")


    

if __name__ == "__main__":
    
    if not os.path.exists(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}"):
        # Create the directory if it does not exist
        os.makedirs(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}")
        
    gr_df = pd.read_csv(rf"V:\StockData\d0_file\IntraAnalysisResults\{today_str}\BrokerFeeSummary\{today_str}_daily_except_monitor_all.csv",encoding='gbk')
    gr_df = gr_df[['产品名','colo','撤单TpCL','撤单TpCrL']].rename(columns = {'产品名':"production"})
    gr_df['TpCrL'] = gr_df['撤单TpCrL'].apply(get_top_cancel_ticker_sum)
    gr_df['TpCrL_sh'] = gr_df[['撤单TpCL','撤单TpCrL']].apply(get_top_cancel_ticker_sum_sh,axis=1)
    gr_df['TpCrL_sz'] = gr_df[['撤单TpCL','撤单TpCrL']].apply(get_top_cancel_ticker_sum_sz,axis=1)
    # df_curr_date 处理保存
    df_curr_date = get_pta_report(today_str,today_str)
    df_curr_date = df_curr_date[['production',"colo"]+Orig_Columns]
    df_curr_date = df_curr_date.rename(columns=Single_Rename_Col)
    df_curr_date["center"] = df_curr_date['colo'].apply(find_colo_shsz)
    df_curr_date_dual = df_curr_date.loc[df_curr_date["production"].isin(DUALCENTER_PRODUCTION)].copy()
    df_curr_date_single = df_curr_date.loc[~df_curr_date["production"].isin(DUALCENTER_PRODUCTION)].copy()
    df_curr_date_dual = get_clean_df(df_curr_date_dual)
    df_curr_date_dual.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_dual_orig.csv",encoding='gbk')
    df_curr_date_single.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_single_orig.csv",encoding='gbk')
    
    # df_10_days 处理保存
    df_10_days = get_pta_report(ten_days_ago,today_str)
    df_10_days = df_10_days[['date','production',"colo"]+Orig_Columns]
    df_10_days = df_10_days.rename(columns=Single_Rename_Col)
    df_10_days["center"] = df_10_days['colo'].apply(find_colo_shsz)
    df_10_days_dual = df_10_days.loc[df_10_days["production"].isin(DUALCENTER_PRODUCTION)].copy()
    df_10_days_single = df_10_days.loc[~df_10_days["production"].isin(DUALCENTER_PRODUCTION)].copy()
    df_10_days_dual = get_clean_df(df_10_days_dual)
    df_10_days_dual.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_{ten_days_ago}_dual_orig.csv",encoding='gbk')
    df_10_days_single.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_{ten_days_ago}_single_orig.csv",encoding='gbk')

    # 计算10 day mean,std,和当日对比平均偏差百分比
    # 筛选自营并保存中间文件（全部t2o数据）
    df_10_days_single_mean = df_10_days_single.loc[:,["production",'center'] + SINGLE_COLUMNS].groupby(["production","center"]).mean().reset_index()
    df_10_days_single_std = df_10_days_single.loc[:,["production",'center'] + SINGLE_COLUMNS].groupby(["production","center"]).std().reset_index()
    df_10_days_single_change_avg,df_10_days_single_change_count = process_daily_average_change(df_10_days_single, cancel = False, dual=False)
    df_10_days_single_mean = df_10_days_single_mean.merge(df_10_days_single_std,how="left",on=['production','center'],suffixes=("","_std"))
    df_10_days_single_mean = df_10_days_single_mean.merge(df_10_days_single_change_avg,how="left",on=['production','center'],suffixes=("","_avg"))
    df_10_days_single_mean = df_10_days_single_mean.merge(df_10_days_single_change_count,how="left",on=['production','center'],suffixes=("","_count"))
    
    
    df_10_days_dual_mean = df_10_days_dual.loc[:,["production",'center'] + DUAL_COLUMNS].groupby(["production","center"]).mean().reset_index()
    df_10_days_dual_std = df_10_days_dual.loc[:,["production",'center'] + DUAL_COLUMNS].groupby(["production","center"]).std().reset_index()
    df_10_days_dual_change_avg,df_10_days_dual_change_count = process_daily_average_change(df_10_days_dual, cancel = False, dual=True)   
    df_10_days_dual_mean = df_10_days_dual_mean.merge(df_10_days_dual_std,how="left",on=['production','center'],suffixes=("","_std"))
    df_10_days_dual_mean = df_10_days_dual_mean.merge(df_10_days_dual_change_avg,how="left",on=['production','center'],suffixes=("","_avg"))
    df_10_days_dual_mean = df_10_days_dual_mean.merge(df_10_days_dual_change_count,how="left",on=['production','center'],suffixes=("","_count"))
    
    df_total_single = df_curr_date_single.merge(df_10_days_single_mean, how="left", on=['production','center'],suffixes=("","_10日"))
    df_total_dual = df_curr_date_dual.merge(df_10_days_dual_mean, how="left", on=['production','center'],suffixes=("","_10日"))
    check_own_acc(df_total_single)
    check_own_acc(df_total_dual)
    df_total_single = df_total_single.sort_values(by="自营",ascending=False)
    df_total_dual = df_total_dual.sort_values(by="自营",ascending=False)
    df_total_single = find_diff_t2o(df_total_single,single=1).fillna(0)
    df_total_dual = find_diff_t2o(df_total_dual,single = 0).fillna(0)
    df_total_single_own = df_total_single[df_total_single["自营"] == 1]
    df_total_dual_own = df_total_dual[df_total_dual["自营"] == 1]
    # df_total_single = df_total_single[df_total_single["自营"] != 1]
    # df_total_dual = df_total_dual[df_total_dual["自营"] != 1]
    
    df_total_dual.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_dual_result.csv",encoding='gbk')
    df_total_single.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_single_result.csv",encoding='gbk')
    df_total_single_own.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_single_own_result.csv",encoding='gbk')
    df_total_dual_own.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_dual_own_result.csv",encoding='gbk')
    
    # 生成输出格式dataframe 并筛选问题产品
    print("\ncombining columns")
    df_total_single = combine_cols(df_total_single)
    df_total_dual = combine_cols(df_total_dual,single=0)
    df_total_single_own = combine_cols(df_total_single_own)
    df_total_dual_own = combine_cols(df_total_dual_own,single=0)
    df_selected_single = filter_df(df_total_single)
    df_selected_dual = filter_df(df_total_dual,single = 0)
    df_selected_single_own = filter_df(df_total_single_own)
    df_selected_dual_own = filter_df(df_total_dual_own,single = 0)
    df_selected_single['指数'] = df_selected_single["production"].apply(lambda prod: PRODUCTION_2_INDEX[prod])
    df_selected_dual['指数'] = df_selected_dual["production"].apply(lambda prod: PRODUCTION_2_INDEX[prod])
    df_selected_single_own['指数'] = df_selected_single_own["production"].apply(lambda prod: PRODUCTION_2_INDEX[prod])
    df_selected_dual_own['指数'] = df_selected_dual_own["production"].apply(lambda prod: PRODUCTION_2_INDEX[prod])
    df_selected_dual.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_dual_bug.csv",encoding='gbk')
    df_selected_single.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_single_bug.csv",encoding='gbk')
    df_selected_single_own.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_single_own_bug.csv",encoding='gbk')
    df_selected_dual_own.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_dual_own_bug.csv",encoding='gbk')
    wechat_bot_markdown("正在生成pta盘后分析报告")
    # 作图+输出
    colored_single_own = process_df_image(df_selected_single_own)
    dfi.export(colored_single_own,filename = rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_single_own_report.png",dpi=500, fontsize=8,max_cols = -1,max_rows=-1,table_conversion='chrome')
    colored_dual_own = process_df_image(df_selected_dual_own,single=0)
    dfi.export(colored_dual_own,filename = rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_dual_own_report.png",dpi=500, fontsize=8,max_cols = -1,max_rows=-1,table_conversion='chrome')
    colored_single = process_df_image(df_selected_single)
    dfi.export(colored_single,filename = rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_single_report.png",dpi=500, fontsize=8,max_cols = -1,max_rows=-1,table_conversion='chrome')
    df_dual_dict = return_df_dict(df_selected_dual)
    dual_t2o_image_lst = []
    for broker in df_dual_dict.keys():
        for df_dual_colo_df in df_dual_dict[broker]:
            if not os.path.exists(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\dual\{broker}"):
                # Create the directory if it does not exist
                os.makedirs(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\dual\{broker}")
            colo_name = list(df_dual_colo_df['simp_colo'])[0]
            df_color_dual = process_df_image(df_dual_colo_df,single = 0)
            dual_t2o_image_lst.append(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\dual\{broker}\dual_{colo_name}_report.png")
            dfi.export(df_color_dual,filename = rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\dual\{broker}\dual_{colo_name}_report.png",dpi=500, fontsize=8,max_cols = -1,max_rows=-1,table_conversion='chrome')
    pdf_generator = GeneratePDF(today_str, rf"V:\StockData\Trading\owen_code\pta_result\{today_str}")
    pdf_generator.generate_pdf(rf"\{today_str}双中心t2o汇总.pdf", dual_t2o_image_lst)
    wechat_bot_file(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}双中心t2o汇总.pdf")
    
    """撤单分开处理"""
    cancel_curr_date = get_pta_report(today_str,today_str)
    cancel_curr_date = cancel_curr_date[['production','colo','撤单股数比(SZ)', '撤单金额比(SZ)','撤单股数比(SH)', '撤单金额比(SH)','撤单股数比','撤单金额比']]
    cancel_ten_date = get_pta_report(ten_days_ago,today_str)
    cancel_ten_date = cancel_ten_date[['date','production','colo','撤单股数比(SZ)', '撤单金额比(SZ)','撤单股数比(SH)', '撤单金额比(SH)','撤单股数比','撤单金额比']]
    cancel_curr_date.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_cancel_orig.csv")
    cancel_ten_date.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_{ten_days_ago}_cancel_orig.csv",encoding='gbk')
    
    cancel_curr_date["center"] = cancel_curr_date['colo'].apply(find_colo_shsz)
    cancel_ten_date["center"] = cancel_ten_date['colo'].apply(find_colo_shsz)
    df_10_days_cancel_mean = cancel_ten_date.loc[:,["production",'center'] + CANCEL_COLUMNS].groupby(["production","center"]).mean().reset_index()
    df_10_days_cancel_change_avg,df_10_days_cancel_change_count = process_daily_average_change(cancel_ten_date, cancel = True, dual=False)
    df_10_days_cancel_mean = df_10_days_cancel_mean.merge(df_10_days_cancel_change_avg,how="left",on=['production','center'],suffixes=("","_avg"))
    df_10_days_cancel_mean = df_10_days_cancel_mean.merge(df_10_days_cancel_change_count,how="left",on=['production','center'],suffixes=("","_count"))
    df_total_cancel = cancel_curr_date.merge(df_10_days_cancel_mean, how="left", on=['production','center'],suffixes=("","_10日"))
    df_total_cancel = find_diff_cancel(df_total_cancel).fillna(0)
    
    # GR个股撤单数据扣除
    df_total_cancel_gr = df_total_cancel.merge(gr_df,how="left",on=['production','colo']).fillna(0)
    df_total_cancel_gr['撤单股数比(SZ)'] = df_total_cancel_gr['撤单股数比(SZ)'] - df_total_cancel_gr['TpCrL_sz']
    df_total_cancel_gr['撤单金额比(SZ)'] = df_total_cancel_gr['撤单金额比(SZ)'] - df_total_cancel_gr['TpCrL_sz']
    df_total_cancel_gr['撤单股数比(SH)'] = df_total_cancel_gr['撤单股数比(SH)'] - df_total_cancel_gr['TpCrL_sh']
    df_total_cancel_gr['撤单金额比(SH)'] = df_total_cancel_gr['撤单金额比(SH)'] - df_total_cancel_gr['TpCrL_sh']
    df_total_cancel_gr['撤单股数比'] = df_total_cancel_gr['撤单股数比'] - df_total_cancel_gr['TpCrL']
    df_total_cancel_gr['撤单金额比'] = df_total_cancel_gr['撤单金额比'] - df_total_cancel_gr['TpCrL']
    
    check_own_acc(df_total_cancel)
    df_total_cancel_own = df_total_cancel[df_total_cancel['自营'] == 1]
    df_total_cancel = df_total_cancel[df_total_cancel['自营'] != 1]
    
    df_total_cancel = combine_cancel_cols(df_total_cancel)
    df_selected_cancel = filter_cancel_df(df_total_cancel)
    df_selected_cancel['指数'] = df_selected_cancel["production"].apply(lambda prod: PRODUCTION_2_INDEX[prod])
    df_total_cancel_own = combine_cancel_cols(df_total_cancel_own).sort_values(by="production",ascending=True)
    df_color_cancel_own = process_cancel_image(df_total_cancel_own,single = 0)
    dfi.export(df_color_cancel_own,filename = rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\cancel_own_report.png",dpi=500, fontsize=8,max_cols = -1,max_rows=-1,table_conversion='chrome')
    df_selected_cancel.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}_{ten_days_ago}_cancel_bug.csv",encoding='gbk')
    df_cancel_dict = return_df_dict(df_selected_cancel)
    
    check_own_acc(df_total_cancel_gr)
    df_total_cancel_own_gr = df_total_cancel_gr[df_total_cancel_gr['自营'] == 1]
    df_total_cancel_gr = df_total_cancel_gr[df_total_cancel_gr['自营'] != 1]
    
    df_total_cancel_gr = combine_cancel_cols(df_total_cancel_gr)
    df_selected_cancel_gr = filter_cancel_df(df_total_cancel_gr)
    df_selected_cancel_gr['指数'] = df_selected_cancel_gr["production"].apply(lambda prod: PRODUCTION_2_INDEX[prod])
    df_total_cancel_own_gr = combine_cancel_cols(df_total_cancel_own_gr).sort_values(by="production",ascending=True)
    df_total_cancel_own_gr = process_cancel_image(df_total_cancel_own_gr,single = 0)
    if not os.path.exists(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}_gr"):
        # Create the directory if it does not exist
        os.makedirs(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}_gr")
    dfi.export(df_total_cancel_own_gr,filename = rf"V:\StockData\Trading\owen_code\pta_result\{today_str}_gr\cancel_own_report_gr.png",dpi=500, fontsize=8,max_cols = -1,max_rows=-1,table_conversion='chrome')
    df_selected_cancel_gr.to_csv(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}_gr\{today_str}_{ten_days_ago}_cancel_bug_gr.csv",encoding='gbk')
    df_cancel_dict_gr = return_df_dict(df_selected_cancel_gr)
    
    #gr个股撤单数据输出
    gr_cancel_image_lst = []
    for broker in df_cancel_dict_gr.keys():
        for df_cancel_colo_df in df_cancel_dict_gr[broker]:
            if not os.path.exists(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}_gr\cancel"):
                # Create the directory if it does not exist
                os.makedirs(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}_gr\cancel")
            colo_name = list(df_cancel_colo_df['simp_colo'])[0]
            df_color_cancel = process_cancel_image(df_cancel_colo_df,single = 0)
            gr_cancel_image_lst.append(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}_gr\cancel\cancel_{colo_name}_report.png")
            dfi.export(df_color_cancel,filename = rf"V:\StockData\Trading\owen_code\pta_result\{today_str}_gr\cancel\cancel_{colo_name}_report.png",dpi=500, fontsize=8,max_cols = -1,max_rows=-1,table_conversion='chrome')
    
    pdf_generator_gr = GeneratePDF(today_str, rf"V:\StockData\Trading\owen_code\pta_result\{today_str}_gr")
    pdf_generator_gr.generate_pdf(rf"\{today_str}管理撤单异常汇总_gr版本.pdf", gr_cancel_image_lst)
    wechat_bot_file(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}_gr\{today_str}管理撤单异常汇总_gr版本.pdf")
    
    # 原版cancel图片输出
    cancel_image_lst = []
    for broker in df_cancel_dict.keys():
        cancel_message_broker = []
        for df_cancel_colo_df in df_cancel_dict[broker]:
            if not os.path.exists(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\cancel\{broker}"):
                # Create the directory if it does not exist
                os.makedirs(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\cancel\{broker}")
            colo_name = list(df_cancel_colo_df['simp_colo'])[0]
            df_color_cancel = process_cancel_image(df_cancel_colo_df,single = 0)
            cancel_image_lst.append(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\cancel\{broker}\cancel_{colo_name}_report.png")
            dfi.export(df_color_cancel,filename = rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\cancel\{broker}\cancel_{colo_name}_report.png",dpi=500, fontsize=8,max_cols = -1,max_rows=-1,table_conversion='chrome')
    
    pdf_generator2 = GeneratePDF(today_str, rf"V:\StockData\Trading\owen_code\pta_result\{today_str}")
    pdf_generator2.generate_pdf(rf"\{today_str}管理撤单异常汇总.pdf", cancel_image_lst)
    wechat_bot_file(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}\{today_str}管理撤单异常汇总.pdf")
    
    # #测试cancel图片输出
    # err_colo = []
    # for broker in df_cancel_dict.keys():
    #     for df_cancel_colo_df in df_cancel_dict[broker]:
    #         if not os.path.exists(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}_demo\cancel"):
    #             # Create the directory if it does not exist
    #             os.makedirs(rf"V:\StockData\Trading\owen_code\pta_result\{today_str}_demo\cancel")
    #         colo_name = list(df_cancel_colo_df['simp_colo'])[0]
    #         if len(df_cancel_colo_df) >= 8:
    #             err_colo.append(colo_name)
    #         df_cancel_colo_df = df_cancel_colo_df.drop_duplicates(subset=CANCEL_COLUMNS + ["production"],keep='first')
    #         df_color_cancel = process_cancel_image(df_cancel_colo_df,single = 0)
    #         dfi.export(df_color_cancel,filename = rf"V:\StockData\Trading\owen_code\pta_result\{today_str}_demo\cancel\cancel_{colo_name}_report.png",dpi=500, fontsize=8,max_cols = -1,max_rows=-1,table_conversion='chrome')
      
    # # 异常输出
    error_output_creation(single = df_selected_single, single_own = df_selected_single_own,
                                      dual_dict = df_dual_dict, dual = df_selected_dual,dual_own = df_selected_dual_own,
                                      cancel_dict = df_cancel_dict, cancel = df_selected_cancel, cancel_own = df_total_cancel_own)
   
    
    """
    撤单10天折线图
    """
    cancel_past_day_graph = get_pta_report(ten_days_ago,today_str)[['date','colo','production','撤单金额比(SZ)', '撤单金额比(SH)']]
    cancel_past_day_graph["center"] = cancel_past_day_graph['colo'].apply(find_colo_shsz)
    for own_prod_subset in [Long_Short_prod,Own_Prod_List]:
        if own_prod_subset == Long_Short_prod:
            label = "LS"
        else:
            label = "own"
        cancel_past_day_prod_type = check_specific_own_acc(cancel_past_day_graph,own_prod_subset)
        graph_past_day_cancel = cancel_past_day_prod_type.loc[cancel_past_day_prod_type["自营"] == 1]
        del graph_past_day_cancel["自营"],graph_past_day_cancel['colo']
        #combine_own_cancel_columns(graph_past_day_cancel)
        # graph_past_day_cancel_sh = graph_past_day_cancel.loc[graph_past_day_cancel['center'] == "sh"]
        # graph_past_day_cancel_sz = graph_past_day_cancel.loc[graph_past_day_cancel['center'] == "sz"]
        plot_cancel_curves(graph_past_day_cancel,today_str,label)


    # plot_t2o_curves("20250520","sh",['DC50E','YX3B2'],['SZ_tick2order(mean)','SH_tick2order(mean)'], 2, 1)
    
    
    